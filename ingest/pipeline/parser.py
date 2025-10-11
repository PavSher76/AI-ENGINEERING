#!/usr/bin/env python3
"""
Парсер документов для модуля "Объекты-аналоги и Архив"
Поддерживает PDF, DOCX, XLSX, IFC, DWG/DXF с OCR fallback
"""

import os
import io
import logging
import asyncio
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
import aiofiles
import hashlib
from datetime import datetime

# PDF обработка
import pdfplumber
import PyPDF2
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image

# Офисные документы
from docx import Document as DocxDocument
import openpyxl
import pandas as pd

# CAD/IFC (опционально)
try:
    import ezdxf
    EZDXF_AVAILABLE = True
except ImportError:
    logging.warning("ezdxf не установлен, DXF файлы не будут обрабатываться")
    EZDXF_AVAILABLE = False

try:
    import ifcopenshell
    IFCOPENSHELL_AVAILABLE = True
except ImportError:
    logging.warning("ifcopenshell не установлен, IFC файлы не будут обрабатываться")
    IFCOPENSHELL_AVAILABLE = False

# OCR
try:
    OCR_AVAILABLE = True
except ImportError:
    logging.warning("OCR библиотеки не установлены")
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Парсер документов с поддержкой множественных форматов и OCR fallback
    """

    def __init__(self, temp_dir: str = "/tmp/parsing"):
        self.temp_dir = temp_dir
        os.makedirs(temp_dir, exist_ok=True)
        
        self.supported_formats = {
            'application/pdf': self._parse_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._parse_docx,
            'application/msword': self._parse_docx,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._parse_xlsx,
            'application/vnd.ms-excel': self._parse_xls,
            'text/plain': self._parse_txt,
            'image/vnd.dxf': self._parse_dxf,
            'application/x-ifc': self._parse_ifc,
        }

    async def parse_document(self, file_path: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Парсит документ и возвращает структурированные данные
        
        Args:
            file_path: Путь к файлу
            metadata: Метаданные документа
            
        Returns:
            Словарь с извлеченными данными
        """
        try:
            # Определяем MIME тип
            mime_type = self._get_mime_type_from_path(file_path)
            
            # Вычисляем хеш файла
            file_hash = await self._calculate_file_hash(file_path)
            file_size = os.path.getsize(file_path)
            
            # Базовые метаданные
            base_metadata = {
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "file_size": file_size,
                "file_hash": file_hash,
                "mime_type": mime_type,
                "parsed_at": datetime.utcnow().isoformat(),
                "parser_version": "1.0.0"
            }
            
            if metadata:
                base_metadata.update(metadata)
            
            # Парсим документ
            parser_func = self.supported_formats.get(mime_type)
            if not parser_func:
                logger.warning(f"Неподдерживаемый формат: {mime_type}")
                return {
                    **base_metadata,
                    "status": "unsupported",
                    "text": "",
                    "tables": [],
                    "page_count": 0,
                    "extraction_method": "none"
                }
            
            # Выполняем парсинг
            result = await parser_func(file_path)
            
            return {
                **base_metadata,
                **result,
                "status": "success"
            }
            
        except Exception as e:
            logger.error(f"Ошибка при парсинге {file_path}: {str(e)}")
            return {
                "file_path": file_path,
                "status": "error",
                "error": str(e),
                "text": "",
                "tables": [],
                "page_count": 0,
                "extraction_method": "failed"
            }

    async def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Парсинг PDF с множественными fallback методами"""
        text = ""
        tables = []
        page_count = 0
        extraction_method = "failed"
        
        try:
            # Метод 1: pdfplumber (лучший для текста и таблиц)
            try:
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    for page_num, page in enumerate(pdf.pages):
                        # Извлекаем текст
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n" + page_text
                        
                        # Извлекаем таблицы
                        page_tables = page.extract_tables()
                        if page_tables:
                            for table_num, table in enumerate(page_tables):
                                tables.append({
                                    "page": page_num + 1,
                                    "table": table_num + 1,
                                    "data": table,
                                    "type": "pdfplumber"
                                })
                
                if text.strip():
                    extraction_method = "pdfplumber"
                    logger.info(f"PDF успешно обработан с pdfplumber: {file_path}")
                    return {
                        "text": text.strip(),
                        "tables": tables,
                        "page_count": page_count,
                        "extraction_method": extraction_method
                    }
                else:
                    logger.warning(f"pdfplumber не извлек текст из: {file_path}")
                    
            except Exception as e:
                logger.warning(f"pdfplumber failed для {file_path}: {e}")

            # Метод 2: PyPDF2 fallback
            try:
                async with aiofiles.open(file_path, 'rb') as file:
                    content = await file.read()
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                    page_count = len(pdf_reader.pages)
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n" + page_text
                
                if text.strip():
                    extraction_method = "pypdf2"
                    logger.info(f"PDF обработан с PyPDF2: {file_path}")
                    return {
                        "text": text.strip(),
                        "tables": tables,
                        "page_count": page_count,
                        "extraction_method": extraction_method
                    }
                else:
                    logger.warning(f"PyPDF2 не извлек текст из: {file_path}")
                    
            except Exception as e:
                logger.warning(f"PyPDF2 failed для {file_path}: {e}")

            # Метод 3: OCR fallback
            if OCR_AVAILABLE:
                try:
                    ocr_text = await self._extract_text_with_ocr(file_path)
                    if ocr_text.strip():
                        extraction_method = "ocr_tesseract"
                        logger.info(f"PDF обработан с OCR: {file_path}")
                        return {
                            "text": ocr_text.strip(),
                            "tables": tables,
                            "page_count": page_count,
                            "extraction_method": extraction_method
                        }
                    else:
                        logger.warning(f"OCR не извлек текст из: {file_path}")
                        
                except Exception as e:
                    logger.warning(f"OCR failed для {file_path}: {e}")
            else:
                logger.warning("OCR недоступен")

            logger.error(f"Все методы парсинга PDF failed для: {file_path}")
            return {
                "text": "",
                "tables": [],
                "page_count": page_count,
                "extraction_method": "failed"
            }

        except Exception as e:
            logger.error(f"Критическая ошибка парсинга PDF {file_path}: {e}")
            raise

    async def _extract_text_with_ocr(self, file_path: str) -> str:
        """OCR извлечение текста из PDF"""
        if not OCR_AVAILABLE:
            raise ImportError("OCR библиотеки недоступны")
        
        extracted_text = ""
        try:
            async with aiofiles.open(file_path, 'rb') as file:
                pdf_content = await file.read()
            
            # Конвертируем PDF в изображения
            images = convert_from_bytes(pdf_content, dpi=300)
            
            for page_num, image in enumerate(images):
                try:
                    # OCR с поддержкой русского и английского
                    custom_config = r'--oem 3 --psm 6 -l rus+eng'
                    page_text = pytesseract.image_to_string(image, config=custom_config)
                    
                    if page_text.strip():
                        extracted_text += f"\n--- Page {page_num + 1} (OCR) ---\n" + page_text
                        
                except Exception as page_error:
                    logger.warning(f"OCR failed для страницы {page_num + 1}: {page_error}")
                    continue
                    
        except Exception as e:
            logger.error(f"OCR ошибка для {file_path}: {e}")
            raise
            
        return extracted_text

    async def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Парсинг DOCX документов"""
        text = ""
        tables = []
        
        try:
            doc = DocxDocument(file_path)
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Извлекаем таблицы
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        "data": table_data,
                        "type": "docx_table"
                    })
            
            logger.info(f"DOCX успешно обработан: {file_path}")
            return {
                "text": text.strip(),
                "tables": tables,
                "extraction_method": "python-docx"
            }
            
        except Exception as e:
            logger.error(f"Ошибка парсинга DOCX {file_path}: {e}")
            raise

    async def _parse_xlsx(self, file_path: str) -> Dict[str, Any]:
        """Парсинг XLSX документов"""
        text = ""
        tables = []
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []
                
                # Извлекаем данные из листа
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    sheet_data.append(row_data)
                
                if sheet_data:
                    tables.append({
                        "sheet_name": sheet_name,
                        "data": sheet_data,
                        "type": "xlsx_sheet"
                    })
                    
                    # Добавляем в текст для поиска
                    text += f"\n--- Sheet: {sheet_name} ---\n"
                    for row_data in sheet_data:
                        text += "\t".join(row_data) + "\n"
            
            logger.info(f"XLSX успешно обработан: {file_path}")
            return {
                "text": text.strip(),
                "tables": tables,
                "extraction_method": "openpyxl"
            }
            
        except Exception as e:
            logger.error(f"Ошибка парсинга XLSX {file_path}: {e}")
            raise

    async def _parse_xls(self, file_path: str) -> Dict[str, Any]:
        """Парсинг XLS документов через pandas"""
        text = ""
        tables = []
        
        try:
            # Читаем все листы
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Конвертируем в список списков
                sheet_data = df.fillna("").astype(str).values.tolist()
                
                if sheet_data:
                    tables.append({
                        "sheet_name": sheet_name,
                        "data": sheet_data,
                        "type": "xls_sheet"
                    })
                    
                    # Добавляем в текст
                    text += f"\n--- Sheet: {sheet_name} ---\n"
                    for row_data in sheet_data:
                        text += "\t".join(row_data) + "\n"
            
            logger.info(f"XLS успешно обработан: {file_path}")
            return {
                "text": text.strip(),
                "tables": tables,
                "extraction_method": "pandas"
            }
            
        except Exception as e:
            logger.error(f"Ошибка парсинга XLS {file_path}: {e}")
            raise

    async def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """Парсинг текстовых файлов"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                text = await file.read()
            
            logger.info(f"TXT успешно обработан: {file_path}")
            return {
                "text": text.strip(),
                "tables": [],
                "extraction_method": "plain_text"
            }
            
        except Exception as e:
            logger.error(f"Ошибка парсинга TXT {file_path}: {e}")
            raise

    async def _parse_dxf(self, file_path: str) -> Dict[str, Any]:
        """Парсинг DXF файлов"""
        if not EZDXF_AVAILABLE:
            logger.warning(f"ezdxf недоступен, пропускаем DXF: {file_path}")
            return {
                "text": "",
                "tables": [],
                "extraction_method": "unsupported"
            }
        
        text_content = []
        try:
            doc = ezdxf.readfile(file_path)
            msp = doc.modelspace()
            
            # Извлекаем текстовые объекты
            for entity in msp:
                if entity.dxftype() in ['TEXT', 'MTEXT']:
                    text_content.append(entity.dxf.text)
                elif entity.dxftype() in ['ATTRIB', 'ATTDEF']:
                    text_content.append(entity.dxf.text)
            
            extracted_text = "\n".join(filter(None, text_content))
            
            logger.info(f"DXF успешно обработан: {file_path}")
            return {
                "text": extracted_text.strip(),
                "tables": [],
                "extraction_method": "ezdxf"
            }
            
        except Exception as e:
            logger.error(f"Ошибка парсинга DXF {file_path}: {e}")
            raise

    async def _parse_ifc(self, file_path: str) -> Dict[str, Any]:
        """Парсинг IFC файлов"""
        if not IFCOPENSHELL_AVAILABLE:
            logger.warning(f"ifcopenshell недоступен, пропускаем IFC: {file_path}")
            return {
                "text": "",
                "tables": [],
                "extraction_method": "unsupported"
            }
        
        text_content = []
        try:
            model = ifcopenshell.open(file_path)
            
            # Извлекаем информацию об объектах
            for entity in model.by_type('IfcRoot'):
                if hasattr(entity, 'Name') and entity.Name:
                    text_content.append(f"Entity: {entity.Name} ({entity.GlobalId})")
                
                # Извлекаем свойства
                for p in entity.IsDefinedBy:
                    if p.is_a('IfcRelDefinesByProperties'):
                        if p.RelatingPropertyDefinition.is_a('IfcPropertySet'):
                            for prop in p.RelatingPropertyDefinition.HasProperties:
                                if prop.is_a('IfcPropertySingleValue'):
                                    value = prop.NominalValue.wrappedValue if prop.NominalValue else 'N/A'
                                    text_content.append(f"  Property: {prop.Name} = {value}")
            
            extracted_text = "\n".join(filter(None, text_content))
            
            logger.info(f"IFC успешно обработан: {file_path}")
            return {
                "text": extracted_text.strip(),
                "tables": [],
                "extraction_method": "ifcopenshell"
            }
            
        except Exception as e:
            logger.error(f"Ошибка парсинга IFC {file_path}: {e}")
            raise

    def _get_mime_type_from_path(self, file_path: str) -> str:
        """Определяет MIME тип по расширению файла"""
        ext = os.path.splitext(file_path)[1].lower()
        
        mime_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.txt': 'text/plain',
            '.dwg': 'image/vnd.dwg',
            '.dxf': 'image/vnd.dxf',
            '.ifc': 'application/x-ifc',
        }
        
        return mime_types.get(ext, 'application/octet-stream')

    async def _calculate_file_hash(self, file_path: str) -> str:
        """Вычисляет SHA-256 хеш файла"""
        hash_sha256 = hashlib.sha256()
        async with aiofiles.open(file_path, 'rb') as f:
            while chunk := await f.read(8192):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()


# Пример использования
async def main():
    """Пример использования парсера"""
    parser = DocumentParser()
    
    # Тестовый файл
    test_file = "test_document.pdf"
    if os.path.exists(test_file):
        result = await parser.parse_document(test_file)
        print(f"Результат парсинга: {result['status']}")
        print(f"Метод извлечения: {result['extraction_method']}")
        print(f"Количество страниц: {result['page_count']}")
        print(f"Длина текста: {len(result['text'])}")
        print(f"Количество таблиц: {len(result['tables'])}")


if __name__ == "__main__":
    asyncio.run(main())
