"""
Сервис для обработки документов
"""

import os
import aiofiles
from typing import Optional, Dict, Any
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import io
import logging
import tempfile

# OCR imports
try:
    from pdf2image import convert_from_bytes
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError as e:
    logging.warning(f"OCR библиотеки не установлены: {e}")
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Класс для извлечения текста из различных форматов документов"""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._extract_from_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_from_docx,
            'application/msword': self._extract_from_docx,
            'text/plain': self._extract_from_txt
        }
    
    async def extract_text(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """
        Извлекает текст из документа
        
        Args:
            file_path: Путь к файлу
            mime_type: MIME тип файла
            
        Returns:
            Словарь с извлеченным текстом и метаданными
        """
        try:
            if mime_type not in self.supported_formats:
                raise ValueError(f"Неподдерживаемый формат файла: {mime_type}")
            
            extractor = self.supported_formats[mime_type]
            result = await extractor(file_path)
            
            logger.info(f"Успешно извлечен текст из {file_path}")
            return result
            
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из {file_path}: {str(e)}")
            raise
    
    async def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Извлекает текст из PDF файла с использованием pdfplumber, PyPDF2 и OCR fallback"""
        text = ""
        tables = []
        page_count = 0
        extraction_method = "unknown"
        
        try:
            # Метод 1: Используем pdfplumber для лучшего извлечения текста
            try:
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    
                    for page_num, page in enumerate(pdf.pages):
                        # Извлекаем текст страницы
                        page_text = page.extract_text()
                        
                        if page_text:
                            text += f"\n--- Страница {page_num + 1} ---\n"
                            text += page_text
                        
                        # Извлекаем таблицы
                        page_tables = page.extract_tables()
                        if page_tables:
                            for table_num, table in enumerate(page_tables):
                                table_data = {
                                    "page": page_num + 1,
                                    "table": table_num + 1,
                                    "data": table
                                }
                                tables.append(table_data)
                
                if text.strip():
                    extraction_method = "pdfplumber"
                    logger.info(f"Успешно извлечен текст из {file_path} с помощью pdfplumber")
                    return {
                        "text": text.strip(),
                        "tables": tables,
                        "page_count": page_count,
                        "extraction_method": extraction_method
                    }
                else:
                    logger.warning(f"pdfplumber не смог извлечь текст из {file_path}")
                    
            except Exception as pdfplumber_error:
                logger.warning(f"Ошибка pdfplumber при обработке {file_path}: {pdfplumber_error}")
            
            # Метод 2: Если pdfplumber не сработал, используем PyPDF2 как fallback
            try:
                logger.info(f"Пробуем PyPDF2 для {file_path}")
                async with aiofiles.open(file_path, 'rb') as file:
                    content = await file.read()
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                    page_count = len(pdf_reader.pages)
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Страница {page_num + 1} ---\n"
                            text += page_text
                
                if text.strip():
                    extraction_method = "pypdf2"
                    logger.info(f"Успешно извлечен текст из {file_path} с помощью PyPDF2")
                    return {
                        "text": text.strip(),
                        "tables": tables,
                        "page_count": page_count,
                        "extraction_method": extraction_method
                    }
                else:
                    logger.warning(f"PyPDF2 не смог извлечь текст из {file_path}")
                    
            except Exception as pypdf2_error:
                logger.warning(f"Ошибка PyPDF2 при обработке {file_path}: {pypdf2_error}")
            
            # Метод 3: Если обычные методы не сработали, используем OCR
            if OCR_AVAILABLE and not text.strip():
                try:
                    logger.info(f"Пробуем OCR для {file_path}")
                    text = await self._extract_text_with_ocr(file_path)
                    if text.strip():
                        extraction_method = "ocr_tesseract"
                        logger.info(f"Успешно извлечен текст из {file_path} с помощью OCR")
                        return {
                            "text": text.strip(),
                            "tables": tables,
                            "page_count": page_count,
                            "extraction_method": extraction_method
                        }
                    else:
                        logger.warning(f"OCR не смог извлечь текст из {file_path}")
                        
                except Exception as ocr_error:
                    logger.warning(f"Ошибка OCR при обработке {file_path}: {ocr_error}")
            elif not OCR_AVAILABLE:
                logger.warning("OCR библиотеки недоступны для fallback")
            
            # Если все методы не сработали
            if not text.strip():
                logger.error(f"Не удалось извлечь текст из {file_path} ни одним из методов")
                return {
                    "text": "",
                    "tables": tables,
                    "page_count": page_count,
                    "extraction_method": "failed"
                }
            
        except Exception as e:
            logger.error(f"Критическая ошибка при обработке PDF {file_path}: {str(e)}")
            raise
    
    async def _extract_text_with_ocr(self, file_path: str) -> str:
        """Извлекает текст из PDF с помощью OCR (Tesseract)"""
        if not OCR_AVAILABLE:
            raise ImportError("OCR библиотеки недоступны")
        
        try:
            # Читаем PDF файл
            async with aiofiles.open(file_path, 'rb') as file:
                pdf_content = await file.read()
            
            # Конвертируем PDF в изображения
            images = convert_from_bytes(pdf_content, dpi=300)
            
            extracted_text = ""
            for page_num, image in enumerate(images):
                try:
                    # Настраиваем Tesseract для русского и английского языков
                    custom_config = r'--oem 3 --psm 6 -l rus+eng'
                    
                    # Извлекаем текст с изображения
                    page_text = pytesseract.image_to_string(image, config=custom_config)
                    
                    if page_text.strip():
                        extracted_text += f"\n--- Страница {page_num + 1} (OCR) ---\n"
                        extracted_text += page_text
                        
                except Exception as page_error:
                    logger.warning(f"Ошибка OCR для страницы {page_num + 1}: {page_error}")
                    continue
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"Ошибка при OCR обработке {file_path}: {str(e)}")
            raise
    
    async def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Извлекает текст из DOCX файла"""
        try:
            text = ""
            tables = []
            
            # Загружаем документ
            doc = DocxDocument(file_path)
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Извлекаем таблицы
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "text": text.strip(),
                "tables": tables,
                "table_count": len(tables),
                "extraction_method": "python-docx"
            }
            
        except Exception as e:
            logger.error(f"Ошибка при обработке DOCX {file_path}: {str(e)}")
            raise
    
    async def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Извлекает текст из текстового файла"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                text = await file.read()
            
            return {
                "text": text.strip(),
                "tables": [],
                "extraction_method": "plain_text"
            }
            
        except Exception as e:
            logger.error(f"Ошибка при обработке TXT {file_path}: {str(e)}")
            raise
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Получает информацию о файле"""
        try:
            stat = os.stat(file_path)
            return {
                "size": stat.st_size,
                "created": stat.st_ctime,
                "modified": stat.st_mtime
            }
        except Exception as e:
            logger.error(f"Ошибка при получении информации о файле {file_path}: {str(e)}")
            return {}
