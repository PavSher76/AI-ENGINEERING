"""
Сервис для парсинга и нормализации документов
"""

import os
import io
import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import aiofiles

# PDF обработка
import pdfplumber
import PyPDF2
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image

# Офисные документы
from docx import Document as DocxDocument
import openpyxl
import pandas as pd

# САПР файлы
try:
    import ifcopenshell
    IFC_AVAILABLE = True
except ImportError:
    IFC_AVAILABLE = False
    logging.warning("ifcopenshell не установлен, IFC файлы не будут обрабатываться")

try:
    import ezdxf
    DXF_AVAILABLE = True
except ImportError:
    DXF_AVAILABLE = False
    logging.warning("ezdxf не установлен, DXF файлы не будут обрабатываться")

from schemas.archive import DocumentMetadata, TextChunk, TableChunk, DrawingChunk, IFCChunk

logger = logging.getLogger(__name__)


class DocumentParser:
    """Парсер документов различных форматов"""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._parse_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._parse_docx,
            'application/msword': self._parse_docx,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._parse_xlsx,
            'application/vnd.ms-excel': self._parse_xlsx,
            'application/ifc': self._parse_ifc,
            'application/dxf': self._parse_dxf,
            'text/plain': self._parse_txt
        }
    
    async def parse_document(self, file_path: str, metadata: DocumentMetadata) -> List[Union[TextChunk, TableChunk, DrawingChunk, IFCChunk]]:
        """
        Парсит документ и возвращает чанки
        
        Args:
            file_path: Путь к файлу
            metadata: Метаданные документа
            
        Returns:
            Список чанков
        """
        try:
            mime_type = self._get_mime_type(file_path)
            
            if mime_type not in self.supported_formats:
                raise ValueError(f"Неподдерживаемый формат файла: {mime_type}")
            
            parser = self.supported_formats[mime_type]
            chunks = await parser(file_path, metadata)
            
            logger.info(f"Документ {file_path} успешно обработан, создано {len(chunks)} чанков")
            return chunks
            
        except Exception as e:
            logger.error(f"Ошибка при парсинге документа {file_path}: {str(e)}")
            raise
    
    async def _parse_pdf(self, file_path: str, metadata: DocumentMetadata) -> List[Union[TextChunk, TableChunk, DrawingChunk]]:
        """Парсит PDF документ"""
        chunks = []
        
        try:
            # Метод 1: pdfplumber для текстовых PDF
            text_chunks = await self._parse_pdf_with_pdfplumber(file_path, metadata)
            chunks.extend(text_chunks)
            
            # Если текст не извлечен, пробуем OCR
            if not text_chunks:
                ocr_chunks = await self._parse_pdf_with_ocr(file_path, metadata)
                chunks.extend(ocr_chunks)
            
            # Если все еще нет чанков, создаем чанк для чертежа
            if not chunks:
                drawing_chunk = await self._create_drawing_chunk(file_path, metadata)
                chunks.append(drawing_chunk)
            
        except Exception as e:
            logger.warning(f"Ошибка при парсинге PDF {file_path}: {str(e)}")
            # Создаем чанк для чертежа как fallback
            drawing_chunk = await self._create_drawing_chunk(file_path, metadata)
            chunks.append(drawing_chunk)
        
        return chunks
    
    async def _parse_pdf_with_pdfplumber(self, file_path: str, metadata: DocumentMetadata) -> List[Union[TextChunk, TableChunk]]:
        """Парсит PDF с помощью pdfplumber"""
        chunks = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Извлекаем текст
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    # Нормализуем текст
                    normalized_text = self._normalize_text(page_text)
                    
                    # Создаем текстовые чанки
                    text_chunks = self._create_text_chunks(
                        normalized_text, metadata, page_num + 1
                    )
                    chunks.extend(text_chunks)
                
                # Извлекаем таблицы
                tables = page.extract_tables()
                for table_num, table in enumerate(tables):
                    if table:
                        table_chunk = self._create_table_chunk(
                            table, metadata, page_num + 1, table_num + 1
                        )
                        chunks.append(table_chunk)
        
        return chunks
    
    async def _parse_pdf_with_ocr(self, file_path: str, metadata: DocumentMetadata) -> List[TextChunk]:
        """Парсит PDF с помощью OCR"""
        chunks = []
        
        try:
            # Конвертируем PDF в изображения
            images = convert_from_bytes(open(file_path, 'rb').read(), dpi=300)
            
            for page_num, image in enumerate(images):
                # Настраиваем Tesseract для русского и английского языков
                custom_config = r'--oem 3 --psm 6 -l rus+eng'
                
                # Извлекаем текст с изображения
                page_text = pytesseract.image_to_string(image, config=custom_config)
                
                if page_text and page_text.strip():
                    # Нормализуем текст
                    normalized_text = self._normalize_text(page_text)
                    
                    # Создаем текстовые чанки
                    text_chunks = self._create_text_chunks(
                        normalized_text, metadata, page_num + 1
                    )
                    chunks.extend(text_chunks)
        
        except Exception as e:
            logger.warning(f"Ошибка OCR для {file_path}: {str(e)}")
        
        return chunks
    
    async def _create_drawing_chunk(self, file_path: str, metadata: DocumentMetadata) -> DrawingChunk:
        """Создает чанк для чертежа"""
        chunk_id = f"drawing_{metadata.doc_no}_{metadata.rev}_{metadata.page or 1}"
        
        # Извлекаем базовую информацию из имени файла
        filename = Path(file_path).stem
        content = f"Чертеж: {filename}\nТип: {metadata.doc_type}\nДисциплина: {metadata.discipline}"
        
        return DrawingChunk(
            chunk_id=chunk_id,
            content=content,
            metadata=metadata,
            chunk_type="drawing",
            preview_path=None,  # Будет создано позже
            extracted_text=None
        )
    
    async def _parse_docx(self, file_path: str, metadata: DocumentMetadata) -> List[TextChunk]:
        """Парсит DOCX документ"""
        chunks = []
        
        try:
            doc = DocxDocument(file_path)
            full_text = []
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            if full_text:
                combined_text = "\n".join(full_text)
                normalized_text = self._normalize_text(combined_text)
                
                # Создаем текстовые чанки
                text_chunks = self._create_text_chunks(normalized_text, metadata)
                chunks.extend(text_chunks)
        
        except Exception as e:
            logger.error(f"Ошибка при парсинге DOCX {file_path}: {str(e)}")
        
        return chunks
    
    async def _parse_xlsx(self, file_path: str, metadata: DocumentMetadata) -> List[TableChunk]:
        """Парсит XLSX документ"""
        chunks = []
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Извлекаем данные из листа
                for row_num, row in enumerate(sheet.iter_rows(values_only=True), 1):
                    if any(cell is not None for cell in row):
                        # Создаем чанк для строки таблицы
                        row_data = {f"col_{i}": str(cell) if cell is not None else "" 
                                   for i, cell in enumerate(row)}
                        
                        # Создаем текстовое представление строки
                        row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        
                        chunk_id = f"table_{metadata.doc_no}_{sheet_name}_row_{row_num}"
                        row_hash = self._calculate_row_hash(row_data)
                        
                        table_chunk = TableChunk(
                            chunk_id=chunk_id,
                            content=row_text,
                            metadata=metadata,
                            chunk_type="table",
                            row_data=row_data,
                            row_hash=row_hash
                        )
                        chunks.append(table_chunk)
        
        except Exception as e:
            logger.error(f"Ошибка при парсинге XLSX {file_path}: {str(e)}")
        
        return chunks
    
    async def _parse_ifc(self, file_path: str, metadata: DocumentMetadata) -> List[IFCChunk]:
        """Парсит IFC документ"""
        chunks = []
        
        if not IFC_AVAILABLE:
            logger.warning("IFC парсинг недоступен, ifcopenshell не установлен")
            return chunks
        
        try:
            ifc_file = ifcopenshell.open(file_path)
            
            # Извлекаем объекты по типам
            object_types = {}
            for entity in ifc_file:
                entity_type = entity.is_a()
                if entity_type not in object_types:
                    object_types[entity_type] = []
                object_types[entity_type].append(entity)
            
            # Создаем чанки для каждого типа объектов
            for entity_type, entities in object_types.items():
                if len(entities) > 0:
                    # Создаем описательный текст для типа объектов
                    content = f"Тип объекта: {entity_type}\nКоличество: {len(entities)}"
                    
                    # Извлекаем свойства первого объекта как пример
                    if entities:
                        sample_entity = entities[0]
                        properties = self._extract_ifc_properties(sample_entity)
                        content += f"\nПример свойств: {properties}"
                    
                    chunk_id = f"ifc_{metadata.doc_no}_{entity_type}"
                    
                    ifc_chunk = IFCChunk(
                        chunk_id=chunk_id,
                        content=content,
                        metadata=metadata,
                        chunk_type="ifc",
                        ifc_type=entity_type,
                        ifc_guid=sample_entity.GlobalId if entities else "",
                        properties=properties
                    )
                    chunks.append(ifc_chunk)
        
        except Exception as e:
            logger.error(f"Ошибка при парсинге IFC {file_path}: {str(e)}")
        
        return chunks
    
    async def _parse_dxf(self, file_path: str, metadata: DocumentMetadata) -> List[DrawingChunk]:
        """Парсит DXF документ"""
        chunks = []
        
        if not DXF_AVAILABLE:
            logger.warning("DXF парсинг недоступен, ezdxf не установлен")
            return chunks
        
        try:
            doc = ezdxf.readfile(file_path)
            msp = doc.modelspace()
            
            # Извлекаем информацию о слоях
            layers = {}
            for entity in msp:
                layer_name = entity.dxf.layer
                if layer_name not in layers:
                    layers[layer_name] = []
                layers[layer_name].append(entity.dxftype())
            
            # Создаем чанк с информацией о слоях
            content = f"DXF файл: {Path(file_path).name}\n"
            content += f"Количество слоев: {len(layers)}\n"
            content += "Слои:\n"
            
            for layer_name, entity_types in layers.items():
                content += f"- {layer_name}: {', '.join(set(entity_types))}\n"
            
            chunk_id = f"dxf_{metadata.doc_no}_{metadata.rev}"
            
            drawing_chunk = DrawingChunk(
                chunk_id=chunk_id,
                content=content,
                metadata=metadata,
                chunk_type="drawing",
                preview_path=None,
                extracted_text=content
            )
            chunks.append(drawing_chunk)
        
        except Exception as e:
            logger.error(f"Ошибка при парсинге DXF {file_path}: {str(e)}")
        
        return chunks
    
    async def _parse_txt(self, file_path: str, metadata: DocumentMetadata) -> List[TextChunk]:
        """Парсит текстовый файл"""
        chunks = []
        
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            if content.strip():
                normalized_text = self._normalize_text(content)
                text_chunks = self._create_text_chunks(normalized_text, metadata)
                chunks.extend(text_chunks)
        
        except Exception as e:
            logger.error(f"Ошибка при парсинге TXT {file_path}: {str(e)}")
        
        return chunks
    
    def _normalize_text(self, text: str) -> str:
        """Нормализует текст"""
        # Убираем лишние пробелы и переносы
        text = " ".join(text.split())
        
        # Исправляем переносы слов
        text = self._fix_word_breaks(text)
        
        # Нормализуем единицы измерения
        text = self._normalize_units(text)
        
        # Нормализуем десятичные разделители
        text = self._normalize_decimal_separators(text)
        
        return text
    
    def _fix_word_breaks(self, text: str) -> str:
        """Исправляет переносы слов"""
        import re
        
        # Исправляем переносы в словах
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
        
        # Исправляем разрывы в числах
        text = re.sub(r'(\d+)\s*\n\s*(\d+)', r'\1\2', text)
        
        # Исправляем разрывы в email
        text = re.sub(r'(\w+)\s*\n\s*@\s*(\w+)', r'\1@\2', text)
        
        return text
    
    def _normalize_units(self, text: str) -> str:
        """Нормализует единицы измерения"""
        import re
        
        # Приводим к SI единицам
        unit_mappings = {
            r'\b(\d+(?:\.\d+)?)\s*мм\b': r'\1 mm',
            r'\b(\d+(?:\.\d+)?)\s*см\b': r'\1 cm',
            r'\b(\d+(?:\.\d+)?)\s*м\b': r'\1 m',
            r'\b(\d+(?:\.\d+)?)\s*кг\b': r'\1 kg',
            r'\b(\d+(?:\.\d+)?)\s*т\b': r'\1 t',
            r'\b(\d+(?:\.\d+)?)\s*°C\b': r'\1 °C',
            r'\b(\d+(?:\.\d+)?)\s*бар\b': r'\1 bar',
            r'\b(\d+(?:\.\d+)?)\s*атм\b': r'\1 atm'
        }
        
        for pattern, replacement in unit_mappings.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        return text
    
    def _normalize_decimal_separators(self, text: str) -> str:
        """Нормализует десятичные разделители"""
        import re
        
        # Заменяем запятую на точку в числах
        text = re.sub(r'(\d+),(\d+)', r'\1.\2', text)
        
        return text
    
    def _create_text_chunks(self, text: str, metadata: DocumentMetadata, page: Optional[int] = None) -> List[TextChunk]:
        """Создает текстовые чанки"""
        chunks = []
        
        # Разбиваем текст на чанки по заголовкам или по размеру
        chunk_size = 800  # токенов
        overlap = 100     # токенов
        
        # Простое разбиение по размеру (в реальном проекте нужно учитывать заголовки)
        words = text.split()
        current_chunk = []
        current_size = 0
        
        for i, word in enumerate(words):
            current_chunk.append(word)
            current_size += 1
            
            if current_size >= chunk_size or i == len(words) - 1:
                chunk_text = " ".join(current_chunk)
                
                chunk_id = f"text_{metadata.doc_no}_{metadata.rev}_{page or 1}_{len(chunks) + 1}"
                
                chunk = TextChunk(
                    chunk_id=chunk_id,
                    content=chunk_text,
                    metadata=metadata,
                    chunk_type="text",
                    token_count=current_size,
                    overlap=overlap if len(chunks) > 0 else 0
                )
                chunks.append(chunk)
                
                # Начинаем новый чанк с перекрытием
                if i < len(words) - 1:
                    current_chunk = current_chunk[-overlap:] if overlap > 0 else []
                    current_size = len(current_chunk)
        
        return chunks
    
    def _create_table_chunk(self, table: List[List[str]], metadata: DocumentMetadata, 
                          page: int, table_num: int) -> TableChunk:
        """Создает чанк таблицы"""
        # Преобразуем таблицу в текстовое представление
        table_text = []
        for row in table:
            if row:
                table_text.append(" | ".join(str(cell) if cell else "" for cell in row))
        
        content = "\n".join(table_text)
        
        # Создаем данные строки (берем первую строку как пример)
        row_data = {}
        if table and table[0]:
            for i, cell in enumerate(table[0]):
                row_data[f"col_{i}"] = str(cell) if cell else ""
        
        chunk_id = f"table_{metadata.doc_no}_{metadata.rev}_{page}_{table_num}"
        row_hash = self._calculate_row_hash(row_data)
        
        return TableChunk(
            chunk_id=chunk_id,
            content=content,
            metadata=metadata,
            chunk_type="table",
            row_data=row_data,
            row_hash=row_hash
        )
    
    def _calculate_row_hash(self, row_data: Dict[str, Any]) -> str:
        """Вычисляет хеш строки таблицы"""
        import hashlib
        import json
        
        # Сортируем ключи для консистентности
        sorted_data = json.dumps(row_data, sort_keys=True)
        return hashlib.md5(sorted_data.encode()).hexdigest()
    
    def _extract_ifc_properties(self, entity) -> Dict[str, Any]:
        """Извлекает свойства IFC объекта"""
        properties = {}
        
        try:
            # Базовые свойства
            if hasattr(entity, 'GlobalId'):
                properties['GlobalId'] = entity.GlobalId
            if hasattr(entity, 'Name'):
                properties['Name'] = entity.Name
            if hasattr(entity, 'Description'):
                properties['Description'] = entity.Description
            
            # Дополнительные свойства из Pset
            if hasattr(entity, 'IsDefinedBy'):
                for rel in entity.IsDefinedBy:
                    if hasattr(rel, 'RelatingPropertyDefinition'):
                        pset = rel.RelatingPropertyDefinition
                        if hasattr(pset, 'HasProperties'):
                            for prop in pset.HasProperties:
                                if hasattr(prop, 'Name') and hasattr(prop, 'NominalValue'):
                                    properties[prop.Name] = prop.NominalValue
        
        except Exception as e:
            logger.warning(f"Ошибка при извлечении свойств IFC: {str(e)}")
        
        return properties
    
    def _get_mime_type(self, file_path: str) -> str:
        """Определяет MIME тип файла"""
        ext = Path(file_path).suffix.lower()
        mime_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.dwg': 'application/dwg',
            '.dxf': 'application/dxf',
            '.ifc': 'application/ifc',
            '.txt': 'text/plain',
            '.rtf': 'application/rtf'
        }
        return mime_types.get(ext, 'application/octet-stream')
