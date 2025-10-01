"""
Сервис для обработки файлов в чате с ИИ
"""

import os
import io
import aiofiles
import magic
import logging
import traceback
from typing import Optional, Dict, Any, List
from pathlib import Path
import asyncio
from datetime import datetime

# Настройка логирования
logger = logging.getLogger(__name__)

# Обработка PDF
import PyPDF2
from pdf2image import convert_from_bytes
import pytesseract

# Обработка DOCX
from docx import Document

# Обработка Excel
import pandas as pd
import openpyxl

# Обработка изображений
from PIL import Image

# OCR
import pytesseract

class FileProcessor:
    """Сервис для обработки различных типов файлов"""
    
    def __init__(self):
        self.max_file_size = 100 * 1024 * 1024  # 100MB
        self.supported_formats = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/vnd.ms-excel': self._process_xls,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_xlsx,
            'text/plain': self._process_txt,
            'text/markdown': self._process_md,
        }
        
        # Настройки OCR
        self.ocr_config = {
            'lang': 'rus+eng',  # Русский и английский
            'config': '--oem 3 --psm 6'
        }
    
    async def process_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Обрабатывает загруженный файл
        
        Args:
            file_content: Содержимое файла
            filename: Имя файла
            
        Returns:
            Словарь с результатами обработки
        """
        logger.info(f"🔄 Начинаем обработку файла: {filename}")
        logger.debug(f"Размер файла: {len(file_content)} байт")
        
        try:
            # Проверяем размер файла
            if len(file_content) > self.max_file_size:
                logger.warning(f"⚠️ Файл {filename} слишком большой: {len(file_content)} > {self.max_file_size}")
                raise ValueError(f"Файл слишком большой. Максимальный размер: {self.max_file_size / (1024*1024):.0f}MB")
            
            # Определяем тип файла
            file_type = magic.from_buffer(file_content, mime=True)
            logger.info(f"📄 Определен тип файла {filename}: {file_type}")
            
            if file_type not in self.supported_formats:
                logger.error(f"❌ Неподдерживаемый формат файла {filename}: {file_type}")
                raise ValueError(f"Неподдерживаемый формат файла: {file_type}")
            
            # Обрабатываем файл
            processor = self.supported_formats[file_type]
            result = await processor(file_content, filename)
            
            return {
                "success": True,
                "filename": filename,
                "file_type": file_type,
                "file_size": len(file_content),
                "processed_at": datetime.now().isoformat(),
                "content": result
            }
            
        except Exception as e:
            logger.error(f"❌ Ошибка обработки файла {filename}: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                "success": False,
                "filename": filename,
                "error": str(e),
                "processed_at": datetime.now().isoformat()
            }
    
    async def _process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Обрабатывает PDF файл"""
        try:
            # Сначала пытаемся извлечь текст напрямую
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_content = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content += f"\n--- Страница {page_num + 1} ---\n"
                        text_content += page_text
                except Exception as e:
                    print(f"Ошибка извлечения текста со страницы {page_num + 1}: {e}")
            
            # Если текст не извлечен или мало текста, используем OCR
            if not text_content.strip() or len(text_content.strip()) < 100:
                print("Применяем OCR для PDF...")
                ocr_text = await self._ocr_pdf(file_content)
                if ocr_text:
                    text_content = ocr_text
            
            return {
                "type": "pdf",
                "text": text_content,
                "pages_count": len(pdf_reader.pages),
                "has_ocr": len(text_content.strip()) > 100 and "OCR" in text_content,
                "metadata": {
                    "title": pdf_reader.metadata.get("/Title", "") if pdf_reader.metadata else "",
                    "author": pdf_reader.metadata.get("/Author", "") if pdf_reader.metadata else "",
                    "creator": pdf_reader.metadata.get("/Creator", "") if pdf_reader.metadata else "",
                }
            }
            
        except Exception as e:
            raise Exception(f"Ошибка обработки PDF: {str(e)}")
    
    async def _ocr_pdf(self, file_content: bytes) -> str:
        """Применяет OCR к PDF файлу"""
        try:
            # Конвертируем PDF в изображения
            images = convert_from_bytes(file_content, dpi=300)
            
            ocr_text = ""
            for i, image in enumerate(images):
                try:
                    # Применяем OCR к изображению
                    page_text = pytesseract.image_to_string(
                        image, 
                        lang=self.ocr_config['lang'],
                        config=self.ocr_config['config']
                    )
                    
                    if page_text.strip():
                        ocr_text += f"\n--- Страница {i + 1} (OCR) ---\n"
                        ocr_text += page_text
                        
                except Exception as e:
                    print(f"Ошибка OCR для страницы {i + 1}: {e}")
                    continue
            
            return ocr_text
            
        except Exception as e:
            print(f"Ошибка OCR: {e}")
            return ""
    
    async def _process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Обрабатывает DOCX файл"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            # Извлекаем текст
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Извлекаем таблицы
            tables_content = ""
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_content += " | ".join(row_text) + "\n"
            
            return {
                "type": "docx",
                "text": text_content,
                "tables": tables_content,
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables),
                "metadata": {
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "created": doc.core_properties.created.isoformat() if doc.core_properties.created else "",
                }
            }
            
        except Exception as e:
            raise Exception(f"Ошибка обработки DOCX: {str(e)}")
    
    async def _process_xls(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Обрабатывает XLS файл"""
        try:
            # Используем pandas для чтения XLS
            df = pd.read_excel(io.BytesIO(file_content), sheet_name=None)
            
            content = {}
            for sheet_name, sheet_df in df.items():
                # Конвертируем DataFrame в текст
                sheet_text = sheet_df.to_string(index=False)
                content[sheet_name] = {
                    "text": sheet_text,
                    "rows": len(sheet_df),
                    "columns": len(sheet_df.columns),
                    "columns_list": list(sheet_df.columns)
                }
            
            return {
                "type": "xls",
                "sheets": content,
                "sheets_count": len(df),
                "sheets_names": list(df.keys())
            }
            
        except Exception as e:
            raise Exception(f"Ошибка обработки XLS: {str(e)}")
    
    async def _process_xlsx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Обрабатывает XLSX файл"""
        try:
            # Используем pandas для чтения XLSX
            df = pd.read_excel(io.BytesIO(file_content), sheet_name=None)
            
            content = {}
            for sheet_name, sheet_df in df.items():
                # Конвертируем DataFrame в текст
                sheet_text = sheet_df.to_string(index=False)
                content[sheet_name] = {
                    "text": sheet_text,
                    "rows": len(sheet_df),
                    "columns": len(sheet_df.columns),
                    "columns_list": list(sheet_df.columns)
                }
            
            return {
                "type": "xlsx",
                "sheets": content,
                "sheets_count": len(df),
                "sheets_names": list(df.keys())
            }
            
        except Exception as e:
            raise Exception(f"Ошибка обработки XLSX: {str(e)}")
    
    async def _process_txt(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Обрабатывает TXT файл"""
        try:
            # Пытаемся определить кодировку
            text_content = file_content.decode('utf-8', errors='ignore')
            
            return {
                "type": "txt",
                "text": text_content,
                "lines_count": len(text_content.split('\n')),
                "chars_count": len(text_content)
            }
            
        except Exception as e:
            raise Exception(f"Ошибка обработки TXT: {str(e)}")
    
    async def _process_md(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Обрабатывает Markdown файл"""
        try:
            text_content = file_content.decode('utf-8', errors='ignore')
            
            return {
                "type": "markdown",
                "text": text_content,
                "lines_count": len(text_content.split('\n')),
                "chars_count": len(text_content)
            }
            
        except Exception as e:
            raise Exception(f"Ошибка обработки Markdown: {str(e)}")
    
    def get_supported_formats(self) -> List[str]:
        """Возвращает список поддерживаемых форматов"""
        return list(self.supported_formats.keys())
    
    def get_max_file_size(self) -> int:
        """Возвращает максимальный размер файла"""
        return self.max_file_size
