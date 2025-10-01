"""
Сервис для экспорта результатов чата в документы
"""

import io
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path

# Экспорт в DOCX
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Экспорт в PDF
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

class DocumentExporter:
    """Сервис для экспорта результатов чата в документы"""
    
    def __init__(self):
        self.setup_fonts()
    
    def setup_fonts(self):
        """Настройка шрифтов для поддержки кириллицы"""
        try:
            # Регистрируем шрифты для поддержки кириллицы
            # В production нужно будет добавить файлы шрифтов
            self.register_cyrillic_fonts()
        except Exception as e:
            print(f"Предупреждение: Не удалось зарегистрировать кириллические шрифты: {e}")
    
    def register_cyrillic_fonts(self):
        """Регистрирует кириллические шрифты"""
        # Попытка зарегистрировать системные шрифты
        try:
            # Для Linux/Windows
            pdfmetrics.registerFont(TTFont('DejaVuSans', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
            pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
        except:
            try:
                # Для macOS
                pdfmetrics.registerFont(TTFont('Arial', '/System/Library/Fonts/Arial.ttf'))
                pdfmetrics.registerFont(TTFont('Arial-Bold', '/System/Library/Fonts/Arial Bold.ttf'))
            except:
                # Fallback на встроенные шрифты
                pass
    
    async def export_to_docx(
        self, 
        chat_data: Dict[str, Any], 
        filename: Optional[str] = None
    ) -> bytes:
        """
        Экспортирует данные чата в DOCX
        
        Args:
            chat_data: Данные чата для экспорта
            filename: Имя файла (опционально)
            
        Returns:
            Байты DOCX файла
        """
        try:
            doc = Document()
            
            # Настройка стилей
            self._setup_docx_styles(doc)
            
            # Заголовок документа
            title = doc.add_heading('Результат работы ИИ', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Метаданные
            doc.add_paragraph(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            doc.add_paragraph(f"Тема: {chat_data.get('topic', 'Чат с ИИ')}")
            
            # Разделитель
            doc.add_paragraph("=" * 50)
            
            # Сообщения чата
            messages = chat_data.get('messages', [])
            for i, message in enumerate(messages):
                self._add_message_to_docx(doc, message, i + 1)
            
            # Файлы (если есть)
            files = chat_data.get('files', [])
            if files:
                doc.add_heading('Прикрепленные файлы', level=1)
                for file_info in files:
                    self._add_file_info_to_docx(doc, file_info)
            
            # Настройки LLM (если есть)
            settings = chat_data.get('llm_settings', {})
            if settings:
                doc.add_heading('Настройки LLM', level=1)
                self._add_settings_to_docx(doc, settings)
            
            # Сохраняем в байты
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return doc_bytes.getvalue()
            
        except Exception as e:
            raise Exception(f"Ошибка экспорта в DOCX: {str(e)}")
    
    async def export_to_pdf(
        self, 
        chat_data: Dict[str, Any], 
        filename: Optional[str] = None
    ) -> bytes:
        """
        Экспортирует данные чата в PDF
        
        Args:
            chat_data: Данные чата для экспорта
            filename: Имя файла (опционально)
            
        Returns:
            Байты PDF файла
        """
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            
            # Настройка стилей
            styles = self._setup_pdf_styles()
            
            # Содержимое документа
            story = []
            
            # Заголовок
            title = Paragraph("Результат работы ИИ", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))
            
            # Метаданные
            meta_data = [
                f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
                f"Тема: {chat_data.get('topic', 'Чат с ИИ')}"
            ]
            
            for meta in meta_data:
                story.append(Paragraph(meta, styles['Normal']))
            
            story.append(Spacer(1, 12))
            story.append(Paragraph("=" * 50, styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Сообщения чата
            messages = chat_data.get('messages', [])
            for i, message in enumerate(messages):
                self._add_message_to_pdf(story, message, i + 1, styles)
            
            # Файлы (если есть)
            files = chat_data.get('files', [])
            if files:
                story.append(Paragraph("Прикрепленные файлы", styles['Heading1']))
                story.append(Spacer(1, 12))
                
                for file_info in files:
                    self._add_file_info_to_pdf(story, file_info, styles)
            
            # Настройки LLM (если есть)
            settings = chat_data.get('llm_settings', {})
            if settings:
                story.append(Paragraph("Настройки LLM", styles['Heading1']))
                story.append(Spacer(1, 12))
                self._add_settings_to_pdf(story, settings, styles)
            
            # Создаем PDF
            doc.build(story)
            buffer.seek(0)
            
            return buffer.getvalue()
            
        except Exception as e:
            raise Exception(f"Ошибка экспорта в PDF: {str(e)}")
    
    def _setup_docx_styles(self, doc: Document):
        """Настраивает стили для DOCX"""
        try:
            # Стиль для пользовательских сообщений
            user_style = doc.styles.add_style('UserMessage', WD_STYLE_TYPE.PARAGRAPH)
            user_style.font.name = 'Arial'
            user_style.font.size = Pt(10)
            user_style.paragraph_format.left_indent = Inches(0.5)
            
            # Стиль для сообщений ИИ
            ai_style = doc.styles.add_style('AIMessage', WD_STYLE_TYPE.PARAGRAPH)
            ai_style.font.name = 'Arial'
            ai_style.font.size = Pt(10)
            ai_style.paragraph_format.left_indent = Inches(0.5)
            ai_style.font.italic = True
            
        except Exception as e:
            print(f"Предупреждение: Не удалось настроить стили DOCX: {e}")
    
    def _setup_pdf_styles(self) -> Dict[str, Any]:
        """Настраивает стили для PDF"""
        styles = getSampleStyleSheet()
        
        # Стиль заголовка
        styles.add(ParagraphStyle(
            name='Title',
            parent=styles['Title'],
            fontName='Helvetica-Bold',
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=12
        ))
        
        # Стиль для пользовательских сообщений
        styles.add(ParagraphStyle(
            name='UserMessage',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leftIndent=20,
            spaceAfter=6
        ))
        
        # Стиль для сообщений ИИ
        styles.add(ParagraphStyle(
            name='AIMessage',
            parent=styles['Normal'],
            fontName='Helvetica-Oblique',
            fontSize=10,
            leftIndent=20,
            spaceAfter=6
        ))
        
        return styles
    
    def _add_message_to_docx(self, doc: Document, message: Dict[str, Any], index: int):
        """Добавляет сообщение в DOCX документ"""
        role = message.get('role', 'user')
        content = message.get('content', '')
        timestamp = message.get('timestamp', '')
        
        # Заголовок сообщения
        if role == 'user':
            heading = doc.add_heading(f"Пользователь ({timestamp})", level=2)
        else:
            heading = doc.add_heading(f"ИИ ({timestamp})", level=2)
        
        # Содержимое сообщения
        content_para = doc.add_paragraph(content)
        
        # Применяем стиль
        if role == 'user':
            content_para.style = 'UserMessage'
        else:
            content_para.style = 'AIMessage'
        
        # Разделитель
        doc.add_paragraph("-" * 30)
    
    def _add_message_to_pdf(self, story: List, message: Dict[str, Any], index: int, styles: Dict[str, Any]):
        """Добавляет сообщение в PDF документ"""
        role = message.get('role', 'user')
        content = message.get('content', '')
        timestamp = message.get('timestamp', '')
        
        # Заголовок сообщения
        if role == 'user':
            heading_text = f"Пользователь ({timestamp})"
        else:
            heading_text = f"ИИ ({timestamp})"
        
        story.append(Paragraph(heading_text, styles['Heading2']))
        
        # Содержимое сообщения
        if role == 'user':
            story.append(Paragraph(content, styles['UserMessage']))
        else:
            story.append(Paragraph(content, styles['AIMessage']))
        
        story.append(Spacer(1, 6))
        story.append(Paragraph("-" * 30, styles['Normal']))
        story.append(Spacer(1, 6))
    
    def _add_file_info_to_docx(self, doc: Document, file_info: Dict[str, Any]):
        """Добавляет информацию о файле в DOCX"""
        filename = file_info.get('filename', 'Неизвестный файл')
        file_type = file_info.get('file_type', '')
        file_size = file_info.get('file_size', 0)
        
        doc.add_paragraph(f"Файл: {filename}")
        doc.add_paragraph(f"Тип: {file_type}")
        doc.add_paragraph(f"Размер: {file_size / 1024:.1f} KB")
        
        # Если есть извлеченный текст
        content = file_info.get('content', {})
        if content.get('text'):
            doc.add_paragraph("Содержимое:")
            doc.add_paragraph(content['text'][:500] + "..." if len(content['text']) > 500 else content['text'])
    
    def _add_file_info_to_pdf(self, story: List, file_info: Dict[str, Any], styles: Dict[str, Any]):
        """Добавляет информацию о файле в PDF"""
        filename = file_info.get('filename', 'Неизвестный файл')
        file_type = file_info.get('file_type', '')
        file_size = file_info.get('file_size', 0)
        
        story.append(Paragraph(f"Файл: {filename}", styles['Normal']))
        story.append(Paragraph(f"Тип: {file_type}", styles['Normal']))
        story.append(Paragraph(f"Размер: {file_size / 1024:.1f} KB", styles['Normal']))
        
        # Если есть извлеченный текст
        content = file_info.get('content', {})
        if content.get('text'):
            story.append(Paragraph("Содержимое:", styles['Normal']))
            text = content['text'][:500] + "..." if len(content['text']) > 500 else content['text']
            story.append(Paragraph(text, styles['Normal']))
        
        story.append(Spacer(1, 6))
    
    def _add_settings_to_docx(self, doc: Document, settings: Dict[str, Any]):
        """Добавляет настройки LLM в DOCX"""
        for key, value in settings.items():
            doc.add_paragraph(f"{key}: {value}")
    
    def _add_settings_to_pdf(self, story: List, settings: Dict[str, Any], styles: Dict[str, Any]):
        """Добавляет настройки LLM в PDF"""
        for key, value in settings.items():
            story.append(Paragraph(f"{key}: {value}", styles['Normal']))
        story.append(Spacer(1, 6))
